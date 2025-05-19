from docx import Document
import os

def extract_text_from_docx(docx_file):
    """Extrae el texto de un archivo DOCX y lo devuelve como string."""
    doc = Document(docx_file)
    full_text = []
    
    # Extraer texto de los párrafos
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extraer texto de las tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

# Lista de archivos DOCX a procesar
docx_files = [
    "White-paper_Continuous-Measurement-CNT-91_No-1_0711_Rev-01.docx",
    "CNT-9X_Programmers_Handbook.docx"
]

# Procesar cada archivo
for docx_file in docx_files:
    if os.path.exists(docx_file):
        try:
            # Extraer el texto
            text = extract_text_from_docx(docx_file)
            
            # Crear nombre del archivo de salida
            txt_file = docx_file.replace('.docx', '.txt')
            
            # Guardar el texto en un archivo
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(text)
            
            print(f"Extraído exitosamente: {docx_file} -> {txt_file}")
        except Exception as e:
            print(f"Error al procesar {docx_file}: {str(e)}")
    else:
        print(f"Archivo no encontrado: {docx_file}") 